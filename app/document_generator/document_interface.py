from abc import ABC, abstractmethod

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.document_generator.text_data import TextData


class IDocument(ABC):
    def __init__(self, ownership, investor):
        self.doc = Document()
        self.ownership = ownership  # ownership object encapsulates name, address, and details
        self.investor = investor


    def add_ownership_data(self):
        pass

    @abstractmethod
    def add_personal_info(self):
        pass

    @abstractmethod
    def add_investor_info(self):
        pass

    @abstractmethod
    def add_details_of_operations(self):
        pass

    @abstractmethod
    def add_agreement_details(self):
        pass

    @abstractmethod
    def add_agreement_couse(self):
        pass

    @abstractmethod
    def add_footer(self0):
        pass

    @abstractmethod
    def add_signature(self):
        pass

    def generate(self, filename):
        self._add_header_section()
        self._add_title_section()
        self.add_personal_info()
        self.add_investor_info()
        self.add_details_of_operations()
        self.add_agreement_details()
        self.add_agreement_couse()
        self.add_footer()
        self.add_signature()
        self.doc.save(filename)

    def _add_header_section(self):
        self._add_right_aligned_paragraph("……………… dnia ..…………….")


    def _add_title_section(self):
        self._add_empty_lines(1)
        self._add_centered_paragraph("OŚWIADCZENIE", bold=True)

    def _add_paragraph(self, text):
        return self.doc.add_paragraph(text)

    def _add_centered_paragraph(self, text, bold=False):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_right_aligned_paragraph(self, text):
        para = self.doc.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _add_left_aligned_paragraph(self, text):
        para = self.doc.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Fix: should be LEFT for left alignment

    def _add_empty_lines(self, count=1):
        for _ in range(count):
            self.doc.add_paragraph()

    def merge_text_into_one_paragraph(self, text1, text2):
        combined_text = text1.text + " " + text2.text
        self.doc.add_paragraph(combined_text)



    def _add_aligned_paragraph(self, left: TextData = None, center: TextData = None, right: TextData = None):
        """
        Adds a paragraph with left-aligned, center-aligned, and right-aligned text on the same line.
        Each position (left, center, right) is represented by a TextData object, which includes the text and formatting.
        """
        para = self.doc.add_paragraph()

        # Add left-aligned text, if provided
        if left:
            self._add_text_with_formatting(para, left.text, left.bold)
            if center or right:
                para.add_run("\t")  # Add tab after left text if there's more text to align

        # Add center-aligned text, if provided
        if center:
            self._add_tab_stop(para, self.doc.sections[0].page_width // 2)
            self._add_text_with_formatting(para, center.text, center.bold)
            if right:
                para.add_run("\t")  # Add tab after center text if there's right-aligned text

        # Add right-aligned text, if provided
        if right:
            self._add_tab_stop(para, self.doc.sections[0].page_width)
            self._add_text_with_formatting(para, right.text, right.bold)

        # Ensure the paragraph itself has no specific alignment (so the tab stops work)
        para.alignment = None



    def _add_text_with_formatting(self, para, text, bold):
        """
        Helper function to add text to a paragraph with optional bold formatting.
        """
        run = para.add_run(text)
        run.bold = bold

    def _add_tab_stop(self, para, position):
        """
        Helper function to add a tab stop at the given position in the paragraph.
        """
        tab_stops = para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(position)

    def add_bold_paragraph(self, text):
        """
        Adds a paragraph with the entire text in bold.
        """
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True

    def add_italic_paragraph(self, text):
        """
        Adds a paragraph with the entire text in italic.
        """
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.italic = True

    def add_paragraph_with_bold_list(self, base_string, variables):
        paragraph = self._add_paragraph(" ")

        # Format the base_string with the provided variables
        formatted_string = base_string.format(*variables)

        # Split the string into parts and identify bold sections (denoted by the variables)
        parts = base_string.split("{")
        variable_index = 0

        for part in parts:
            if "}" in part:
                # Add the bold part (variable)
                paragraph.add_run(variables[variable_index]).bold = True
                variable_index += 1
                # Add the rest of the text after the placeholder
                remaining_text = part.split("}", 1)[1]  # Get the text after the placeholder
                if remaining_text:
                    paragraph.add_run(remaining_text)
            else:
                # Add any text that comes before the first placeholder
                paragraph.add_run(part)




