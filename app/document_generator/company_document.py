
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.document_generator.document_interface import IDocument


class CompanyDocument(IDocument):
    def add_personal_info(self):
        self._add_empty_lines(1)

        self.add_paragraph_with_bold_list(
            "Ja niżej podpisany(a) {0} legitymujący(a) {1} siedziba firmy {2}",
            [self.ownership.get_name(), self.ownership.get_details(), self.ownership.get_address()])

    def add_investor_info(self):
        self.add_paragraph_with_bold_list(
            "w związku z inwestycją planowaną przez Inwestora {0} z siedzibą {1}, polegającą na budowie obiektu elektroenergetycznego pod nazwą: ",
            [self.investor.get_name(), self.investor.get_address()])

    def add_details_of_operations(self):
        self._add_centered_paragraph(self.investor.get_name_of_work(), bold=True)

        self.add_paragraph_with_bold_list(
            "oświadczam, że jako właściciel nieruchomości wyrażam zgodę na dysponowanie nieruchomością do celów"
            " budowlanych, dla działek oznaczonych w ewidencji gruntów i budynków jako działki nr {0} w obrębie"
            " ewidencyjnym {1} wpisanych w KW nr {2} prowadzonej przez Sąd Rejonowy w Tucholi.",
            [self.ownership.get_parcel_number(), self.ownership.get_circle(), self.ownership.get_land_register()])

        self._add_left_aligned_paragraph("Jednocześnie oświadczam, że: ")
        para1 = self.doc.add_paragraph()
        para1.style = 'ListBullet'
        run1 = para1.add_run(
            "zapoznałem(am) się z zakresem ww. inwestycji i wyrażam zgodę na realizację całego zakresu prac na odcinku przebiegającym przez moją nieruchomość:")

        para2 = self.doc.add_paragraph()
        para2.style = 'ListBullet'
        run2 = para2.add_run(self.investor.get_type_of_work()).bold = True

        para3 = self.doc.add_paragraph()
        para3.style = 'ListBullet'
        run3 = para3.add_run(
            "wyrażam zgodę na wejście na nieruchomość w celu wykonania ww. prac oraz ewentualnych prac demontażowych,")

    def add_agreement_details(self):
        """
        Adds a paragraph with the agreement details, where specific parts are bold and italic.
        """
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left-align the paragraph

        self._add_italic_run(para,
                             "Równocześnie wyrażam zgodę na ustanowienie nieodpłatnej służebności przesyłu na rzecz ")
        self._add_bold_and_italic_run(para, self.investor.get_name())
        self._add_italic_run(para, " na dz. nr ")
        self._add_bold_and_italic_run(para, self.ownership.get_parcel_number())
        self._add_italic_run(para, " obręb ")
        self._add_bold_and_italic_run(para, self.ownership.get_circle())
        self._add_italic_run(para, " w m. ")
        self._add_bold_and_italic_run(para, self.ownership.reg_unit)
        self._add_italic_run(para, ", w zakresie sieci elektroenergetycznej.")

        self._add_left_aligned_paragraph(
            f"Koszty związane z ustanowieniem i ujawnieniem służebności przesyłu pokryje {self.investor.get_name()}")

    def add_agreement_couse(self):
        self._add_left_aligned_paragraph(
            f"Odszkodowanie za ewentualnie powstałe szkody w wyniku realizowanych robót pokrywa"
            f" wykonawca działający w imieniu i na rzecz {self.investor.get_name()} na podstawie protokołów"
            f" oszacowania szkód sporządzonych komisyjnie przy udziale wykonawcy robót, inspektora"
            f" nadzoru i osoby bezpośrednio poszkodowanej. Jednocześnie wykonawca robót zobowiązany"
            f" jest do przywrócenia terenu do stanu pierwotnego.")

        self._add_left_aligned_paragraph(
            "Świadomy(a) odpowiedzialności karnej za prawdziwość wskazanych wyżej danych na zasadzie "
            " art. 233 Kodeksu karnego, potwierdzam ich prawdziwość przez złożenie własnoręcznego "
            " podpisu na niniejszym oświadczeniu.")

    def add_footer(self):
        self.add_bold_paragraph(
            "Załącznikiem do niniejszego oświadczenia jest mapa sytuacyjno-wysokościowa z oznaczonym"
            " przebiegiem projektowanej infrastruktury elektroenergetycznej, która winna być podpisana"
            " przez właściciela nieruchomości.")

        self._add_empty_lines(1)

        self.add_bold_paragraph("Tel. .........………………")

        self._add_left_aligned_paragraph(
            f"Oświadczam, że zapoznałem się z informacją dotyczącą przetwarzania moich danych osobowych"
            f" przez {self.investor.get_name()} przekazaną mi wraz z  niniejszym oświadczeniem.")

    def add_signature(self):
        self._add_paragraph(
            "Świadomy(a) odpowiedzialności karnej za prawdziwość wskazanych wyżej danych na zasadzie art. 233 Kodeksu karnego,"
            " potwierdzam ich prawdziwość przez złożenie własnoręcznego podpisu na niniejszym oświadczeniu."
        )

        self._add_left_aligned_paragraph(".........……………… dnia  ........................ ")
        self._add_right_aligned_paragraph("....................................................")

    def _add_italic_run(self, para, text):
        """
        Helper function to add an italic run to a paragraph.
        """
        run = para.add_run(text)
        run.italic = True

    def _add_bold_and_italic_run(self, para, text):
        """
        Helper function to add a bold and italic run to a paragraph.
        """
        run = para.add_run(text)
        run.bold = True
        run.italic = True
