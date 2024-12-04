from docx import Document



class DocumentGenerator:
    def __init__(self, ownership):
        # Initialize the document object
        self.doc = Document()
        self.ownership = ownership

    def add_paragraph(self, text, bold=False, underline=False, alignment=None):
        # Create and add a paragraph with optional styling
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        if bold:
            run.bold = True
        if underline:
            run.underline = True
        if alignment:
            para.alignment = alignment
        return para

    def add_empty_lines(self, count=1):
        # Add empty lines
        for _ in range(count):
            self.doc.add_paragraph()


